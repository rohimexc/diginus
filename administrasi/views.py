from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.utils import timezone
from .models import *
from django.db.models import Sum
from decimal import Decimal
from django.db import transaction
import io
from docx import Document


# --- HELPER: GENERIC DELETE ---
@require_POST
def api_generic_delete(request, model_name, pk):
    """Satu fungsi untuk menghapus semua jenis data secara aman"""
    mapping = {
        'customer': Customer,
        'doc': Correspondence,
        'type': DocumentType,
        'expense': Expense,
        'invoice': Invoice,
    }
    model = mapping.get(model_name)
    if not model:
        return JsonResponse({'status': 'error', 'message': 'Model tidak ditemukan'}, status=400)
    
    obj = get_object_or_404(model, pk=pk)
    obj.delete()
    return JsonResponse({'status': 'success', 'message': 'Data berhasil dihapus'})


# --- CUSTOMER MANAGEMENT ---
def customer_list(request):
    return render(request, 'administrasi/customer_list.html')

def api_customer_data(request):
    customers = list(Customer.objects.values().order_by('-created_at'))
    return JsonResponse({'data': customers})

@require_POST
def api_customer_upsert(request):
    pk = request.POST.get('id')
    data = {
        'name': request.POST.get('name'),
        'company': request.POST.get('company'),
        'email': request.POST.get('email'),
        'whatsapp': request.POST.get('whatsapp'),
        'address': request.POST.get('address'),
    }
    try:
        if pk:
            Customer.objects.filter(pk=pk).update(**data)
            msg = "Data pelanggan diperbarui"
        else:
            Customer.objects.create(**data)
            msg = "Pelanggan baru ditambahkan"
        return JsonResponse({'status': 'success', 'message': msg})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)


# --- DOCUMENT TYPE MANAGEMENT ---
def document_type_list(request):
    return render(request, 'administrasi/document_type_list.html')

def api_document_type_data(request):
    types = list(DocumentType.objects.values('id', 'name', 'code', 'template_docx'))
    return JsonResponse({'data': types})

@require_POST
def api_document_type_upsert(request):
    pk = request.POST.get('id')
    name = request.POST.get('name')
    code = request.POST.get('code').upper()
    template_file = request.FILES.get('template_docx')

    try:
        if pk:
            dt = get_object_or_404(DocumentType, pk=pk)
            dt.name = name
            dt.code = code
            if template_file: dt.template_docx = template_file
            dt.save()
            msg = "Jenis surat diperbarui"
        else:
            DocumentType.objects.create(name=name, code=code, template_docx=template_file)
            msg = "Jenis surat ditambahkan"
        return JsonResponse({'status': 'success', 'message': msg})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)


# --- CORRESPONDENCE MANAGEMENT ---
def correspondence_list(request):
    context = {
        'customers': Customer.objects.all(),
        'doc_types': DocumentType.objects.all()
    }
    return render(request, 'administrasi/correspondence_list.html', context)

def api_correspondence_data(request):
    docs = Correspondence.objects.select_related('customer', 'doc_type', 'invoice_detail').all().order_by('-created_at')
    data = []
    for d in docs:
        data.append({
            'id': d.id,
            'number': d.formatted_number,
            'customer': d.customer.name,
            'type': d.doc_type.name,
            'invoice_id': d.invoice_detail.id if hasattr(d, 'invoice_detail') else None,
            'subject': d.subject,
            'pdf_url': d.file_pdf.url if d.file_pdf else None,
            'docx_url': d.file_docx.url if d.file_docx else None,
            'created_at': d.created_at.strftime('%d/%m/%Y')
        })
    return JsonResponse({'data': data})

@require_POST
def api_correspondence_upsert(request):
    """Handle Create & Update Surat"""
    pk = request.POST.get('id')
    customer_id = request.POST.get('customer_id')
    doc_type_id = request.POST.get('doc_type_id')
    subject = request.POST.get('subject')

    try:
        if pk:
            doc = get_object_or_404(Correspondence, pk=pk)
            doc.customer_id = customer_id
            doc.doc_type_id = doc_type_id
            doc.subject = subject
            doc.save()
            msg = "Data surat diperbarui"
        else:
            doc = Correspondence.objects.create(
                customer_id=customer_id, doc_type_id=doc_type_id, subject=subject
            )
            # Otomatis buat baris Invoice jika jenis surat adalah 'INV' atau 'INVOICE'
            if doc.doc_type.code == 'INV':
                Invoice.objects.get_or_create(correspondence=doc)
            msg = "Surat baru berhasil diterbitkan"
        
        return JsonResponse({'status': 'success', 'message': msg, 'number': doc.formatted_number})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)

def print_docx(request, pk):
    # 1. Ambil data surat berdasarkan ID
    doc_obj = get_object_or_404(Correspondence, pk=pk)
    
    # 2. Ambil path file template dari model DocumentType
    # doc_obj.doc_type.template_docx.path mengambil lokasi file di server
    try:
        template_path = doc_obj.doc_type.template_docx.path
        document = Document(template_path)
    except Exception as e:
        return HttpResponse(f"Template tidak ditemukan: {e}", status=404)

    # 3. Mapping variabel sesuai dengan isi template penawaran Anda
    replacements = {
        '{{number}}': doc_obj.formatted_number, # [cite: 1]
        '{{subject}}': doc_obj.subject,          # [cite: 2]
        '{{company}}': doc_obj.customer.company or doc_obj.customer.name,  
        '{{tanggal_surat}}': doc_obj.created_at.strftime('%d %B %Y'), 
    }

    # 4. Proses penggantian teks di semua paragraf
    for paragraph in document.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                # Menjaga formatting asli template
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # 5. Proses penggantian teks di dalam tabel (jika ada tabel di template)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

    # 6. Kirim file sebagai download stream
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = f"{doc_obj.doc_type.code}_{doc_obj.number}.docx"
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response


# --- INVOICE & FINANCE MANAGEMENT ---
def invoice_list_page(request):
    return render(request, 'administrasi/invoice_list.html')

def api_invoice_data(request):
    # Gunakan prefetch_related agar loading data item lebih cepat
    invoices = Invoice.objects.select_related('correspondence', 'correspondence__customer').prefetch_related('items').all()
    data = []
    
    for inv in invoices:
        # Susun daftar item untuk dikirim ke Modal Edit di Frontend
        items_list = []
        for item in inv.items.all():
            items_list.append({
                'desc': item.description,
                'qty': item.quantity,
                'price': float(item.unit_price),
                'subtotal': float(item.subtotal)
            })
            
        data.append({
            'id': inv.id,
            'number': inv.correspondence.formatted_number,
            'customer': inv.correspondence.customer.name,
            # Gunakan properti yang ada di models.py Anda
            'total_amount': float(inv.grand_total), # Ini nilai setelah PPN 11%
            'paid_amount': float(inv.paid_amount),
            'remaining': float(inv.remaining_balance),
            'is_paid': inv.is_paid,
            'invoice_type': inv.invoice_type,
            'due_date': inv.due_date.strftime('%Y-%m-%d') if inv.due_date else '',
            'items': items_list 
        })
        
    return JsonResponse({'data': data})

@require_POST
def api_invoice_upsert(request):
    pk = request.POST.get('id')
    invoice_type = request.POST.get('invoice_type', 'LUNAS')
    due_date = request.POST.get('due_date')
    paid_amount = request.POST.get('paid_amount') or 0
    
    # Ambil data array dari form (dikirim oleh item_desc[], dll)
    item_descriptions = request.POST.getlist('item_desc[]')
    item_quantities = request.POST.getlist('item_qty[]')
    item_prices = request.POST.getlist('item_price[]')

    try:
        # Gunakan transaction agar jika error, data tidak tersimpan setengah-setengah
        with transaction.atomic():
            # 1. Cari atau buat Invoice Header
            # Note: Invoice biasanya sudah dibuat otomatis saat Correspondence dibuat
            invoice = Invoice.objects.get(pk=pk)
            
            # 2. Update field utama Invoice
            invoice.invoice_type = invoice_type
            invoice.paid_amount = Decimal(paid_amount)
            if due_date:
                invoice.due_date = due_date
            
            # 3. Proses Item (Hapus item lama, lalu buat yang baru)
            # Cara paling aman: Hapus semua item yang ada, lalu insert ulang dari form
            invoice.items.all().delete()
            
            new_items = []
            for i in range(len(item_descriptions)):
                if item_descriptions[i].strip(): # Hanya simpan jika deskripsi tidak kosong
                    new_items.append(InvoiceItem(
                        invoice=invoice,
                        description=item_descriptions[i],
                        quantity=int(item_quantities[i]),
                        unit_price=Decimal(item_prices[i])
                    ))
            
            # Simpan semua item baru sekaligus (bulk create lebih efisien)
            InvoiceItem.objects.bulk_create(new_items)
            
            # 4. Update status Lunas otomatis jika sisa = 0
            # Kita panggil total_amount dari @property yang menghitung dari database terbaru
            if invoice.remaining_balance <= 0:
                invoice.is_paid = True
            else:
                invoice.is_paid = False
                
            invoice.save()

            return JsonResponse({
                'status': 'success', 
                'message': 'Invoice dan item berhasil diperbarui'
            })

    except Invoice.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Data Invoice tidak ditemukan'}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)

def print_invoice(request, pk):
    # Mengambil data invoice beserta itemnya
    invoice = get_object_or_404(Invoice, pk=pk)
    
    context = {
        'invoice': invoice,
    }
    return render(request, 'administrasi/invoice_template.html', context)

# --- EXPENSE & DASHBOARD ---
def finance_dashboard(request):
    return render(request, 'administrasi/finance_dashboard.html')

def api_finance_summary(request):
    month = request.GET.get('month', timezone.now().month)
    year = request.GET.get('year', timezone.now().year)
    report = Expense.get_financial_report(month=int(month), year=int(year))
    data = {k: float(v) for k, v in report.items()}
    return JsonResponse(data)

def api_expense_data(request):
    expenses = list(Expense.objects.values().order_by('-date'))
    return JsonResponse({'data': expenses})

@require_POST
def api_expense_upsert(request):
    pk = request.POST.get('id')
    
    # Ambil data dengan proteksi default value
    category = request.POST.get('category')
    title = request.POST.get('title')
    amount_raw = request.POST.get('amount')
    date_raw = request.POST.get('date')

    # Validasi Dasar: Pastikan data penting ada
    if not all([category, title, amount_raw]):
        return JsonResponse({'status': 'error', 'message': 'Kategori, Judul, dan Nominal harus diisi'}, status=400)

    try:
        data = {
            'category': category,
            'title': title,
            'amount': Decimal(amount_raw),
            'date': date_raw if date_raw else timezone.now().date(),
            'note': request.POST.get('note', ''),
        }
        
        if pk and pk.strip():
            Expense.objects.filter(pk=pk).update(**data)
            msg = "Data diperbarui"
        else:
            Expense.objects.create(**data)
            msg = "Data ditambahkan"
            
        return JsonResponse({'status': 'success', 'message': msg})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)